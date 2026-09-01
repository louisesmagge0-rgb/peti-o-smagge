from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse, HTMLResponse
from pydantic import BaseModel, Field
from docx import Document
from typing import Dict
import os
import uuid


app = FastAPI(
    title="SMAGGE - Gerador de Petições",
    version="1.0.0"
)


PASTA_SAIDA = "arquivos"
os.makedirs(PASTA_SAIDA, exist_ok=True)


class PedidoPeticao(BaseModel):
    tipo_peticao: str
    nome_cliente: str
    substituicoes: Dict[str, str] = Field(default_factory=dict)


def localizar_modelo(tipo: str):
    arquivos = os.listdir(".")
    tipo = tipo.lower()

    if tipo == "juros":
        for arq in arquivos:
            nome = arq.lower()

            if (
                arq.lower().endswith(".docx")
                and "juros" in nome
                and "abusiv" in nome
            ):
                return arq

    if tipo == "descumprimento":
        for arq in arquivos:
            nome = arq.lower()

            if (
                arq.lower().endswith(".docx")
                and "descumprimento" in nome
            ):
                return arq

    return None


def substituir_em_runs(paragrafo, substituicoes):
    if not paragrafo.runs:
        return

    texto = "".join(run.text for run in paragrafo.runs)
    original = texto

    for antigo, novo in substituicoes.items():
        if antigo:
            texto = texto.replace(antigo, str(novo))

    if texto == original:
        return

    paragrafo.runs[0].text = texto

    for run in paragrafo.runs[1:]:
        run.text = ""


def processar_documento(doc, substituicoes):

    for paragrafo in doc.paragraphs:
        substituir_em_runs(paragrafo, substituicoes)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    substituir_em_runs(paragrafo, substituicoes)

    for secao in doc.sections:

        for paragrafo in secao.header.paragraphs:
            substituir_em_runs(paragrafo, substituicoes)

        for tabela in secao.header.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        substituir_em_runs(paragrafo, substituicoes)

        for paragrafo in secao.footer.paragraphs:
            substituir_em_runs(paragrafo, substituicoes)


@app.get("/", response_class=HTMLResponse)
def inicio():

    return """
<!DOCTYPE html>
<html lang="pt-BR">

<head>

<meta charset="UTF-8">

<meta name="viewport" content="width=device-width, initial-scale=1.0">

<title>SMAGGE - Gerador de Petições</title>

<style>

body {
    font-family: Arial, sans-serif;
    background: #f4f4f4;
    margin: 0;
    padding: 0;
}

.topo {
    background: #111827;
    color: white;
    padding: 25px;
    text-align: center;
}

.container {
    background: white;
    max-width: 800px;
    margin: 40px auto;
    padding: 35px;
    border-radius: 8px;
    box-shadow: 0 0 15px rgba(0,0,0,.15);
}

h1 {
    margin: 0;
}

h2 {
    color: #111827;
}

label {
    font-weight: bold;
    display: block;
    margin-top: 18px;
}

input,
select,
textarea {
    width: 100%;
    box-sizing: border-box;
    padding: 12px;
    margin-top: 6px;
    border: 1px solid #ccc;
    border-radius: 5px;
}

textarea {
    height: 250px;
    font-family: monospace;
}

button {
    width: 100%;
    margin-top: 25px;
    padding: 15px;
    border: none;
    background: #111827;
    color: white;
    font-size: 16px;
    font-weight: bold;
    border-radius: 5px;
    cursor: pointer;
}

button:hover {
    background: #374151;
}

#mensagem {
    margin-top: 20px;
    font-weight: bold;
}

.aviso {
    margin-top: 20px;
    padding: 15px;
    background: #f3f4f6;
    border-left: 4px solid #111827;
}

</style>

</head>

<body>

<div class="topo">

<h1>SMAGGE ADVOGADOS ASSOCIADOS</h1>

<p>Gerador de Petições</p>

</div>

<div class="container">

<h2>Gerar Petição</h2>

<label>Tipo de petição</label>

<select id="tipo">

<option value="juros">
Juros Abusivos
</option>

<option value="descumprimento">
Descumprimento Contratual
</option>

</select>

<label>Nome do cliente</label>

<input
id="cliente"
placeholder="Digite o nome completo">

<label>Substituições</label>

<textarea id="substituicoes">
{
    "XXXXX": "DADO DO CLIENTE"
}
</textarea>

<div class="aviso">

Informe no campo acima os textos existentes no modelo e os respectivos textos que devem substituí-los.

</div>

<button onclick="gerar()">
GERAR PETIÇÃO
</button>

<div id="mensagem"></div>

</div>

<script>

async function gerar() {

    const mensagem =
        document.getElementById("mensagem");

    mensagem.innerHTML =
        "Gerando petição...";

    let substituicoes;

    try {

        substituicoes =
            JSON.parse(
                document.getElementById(
                    "substituicoes"
                ).value
            );

    }

    catch (erro) {

        mensagem.innerHTML =
            "Erro: confira o campo de substituições.";

        return;
    }

    const dados = {

        tipo_peticao:
            document.getElementById(
                "tipo"
            ).value,

        nome_cliente:
            document.getElementById(
                "cliente"
            ).value,

        substituicoes:
            substituicoes
    };

    try {

        const resposta =
            await fetch(
                "/gerar",
                {
                    method: "POST",

                    headers: {
                        "Content-Type":
                            "application/json"
                    },

                    body:
                        JSON.stringify(dados)
                }
            );

        const resultado =
            await resposta.json();

        if (!resposta.ok) {

            mensagem.innerHTML =
                "Erro: " +
                (
                    resultado.detail ||
                    "Não foi possível gerar."
                );

            return;
        }

        mensagem.innerHTML =
            "Petição gerada. Iniciando download...";

        window.location.href =
            resultado.download;

    }

    catch (erro) {

        mensagem.innerHTML =
            "Erro de comunicação com o servidor.";

    }

}

</script>

</body>

</html>
"""


@app.get("/status")
def status():
    return {
        "status": "online",
        "sistema": "SMAGGE Gerador de Petições"
    }


@app.post("/gerar")
def gerar_peticao(pedido: PedidoPeticao):

    modelo = localizar_modelo(
        pedido.tipo_peticao
    )

    if not modelo:
        raise HTTPException(
            status_code=404,
            detail="Modelo da petição não encontrado."
        )

    try:
        doc = Document(modelo)

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Erro ao abrir modelo: {str(e)}"
        )

    processar_documento(
        doc,
        pedido.substituicoes
    )

    nome_seguro = "".join(
        c if c.isalnum() else "_"
        for c in pedido.nome_cliente
    )

    codigo = str(uuid.uuid4())

    nome_arquivo = (
        f"{codigo}_PETICAO_{nome_seguro}.docx"
    )

    caminho = os.path.join(
        PASTA_SAIDA,
        nome_arquivo
    )

    try:
        doc.save(caminho)

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Erro ao salvar documento: {str(e)}"
        )

    return {
        "sucesso": True,
        "arquivo_id": codigo,
        "arquivo": nome_arquivo,
        "download": f"/download/{codigo}"
    }


@app.get("/download/{arquivo_id}")
def baixar_peticao(arquivo_id: str):

    for arquivo in os.listdir(PASTA_SAIDA):

        if arquivo.startswith(arquivo_id):

            caminho = os.path.join(
                PASTA_SAIDA,
                arquivo
            )

            return FileResponse(
                caminho,
                filename=arquivo,
                media_type=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                )
            )

    raise HTTPException(
        status_code=404,
        detail="Arquivo não encontrado."
    )
